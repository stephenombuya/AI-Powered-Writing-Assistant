from docx import Document as DocxDocument
from fpdf import FPDF
import os
from typing import Optional

class DocumentExporter:
    @staticmethod
    def export_to_word(content: str, title: str) -> str:
        """Export content to Word document."""
        doc = DocxDocument()
        doc.add_heading(title, 0)
        doc.add_paragraph(content)
        
        filename = f"{title.lower().replace(' ', '_')}.docx"
        doc.save(filename)
        return filename
    
    @staticmethod
    def export_to_pdf(content: str, title: str) -> str:
        """Export content to PDF."""
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        # Add title
        pdf.set_font("Arial", 'B', 16)
        pdf.cell(200, 10, txt=title, ln=1, align='C')
        
        # Add content
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, txt=content)
        
        filename = f"{title.lower().replace(' ', '_')}.pdf"
        pdf.output(filename)
        return filename

